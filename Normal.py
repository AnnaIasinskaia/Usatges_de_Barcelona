#!/usr/bin/env python
# coding: utf-8

# In[203]:


from docx import Document
import docx
from numpy import unravel_index
import re
import numpy as np


# In[204]:


#document = Document("Lex visigothorum.docx")


# In[205]:


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


# In[206]:


text = getText("Lex visigothorum.docx")


# In[207]:


textList = text.split('\n')


# In[226]:


#textList=textList[:4000]


# In[227]:


print(textList[0:10])


# In[228]:



new = []
for x in textList:
    #выкидывает параграфы, в которых содержатся ненужные символы(использовано для примечаний)
    if ('|' in x):
        continue
    if ('HN' in x):
        continue
    if ('PHV' in x):
        continue
    if ('PNV' in x):
        continue
    y=x
    y=y.lower()
    y=re.sub(r'[^a-z ]+', '', y)
    
    if not x:
        continue
    print(y)
    new.append(y)


# In[229]:


myWordList=[]
for x in new:
    myWordList+=x.split()
print (myWordList)


# In[230]:


len(myWordList)


# In[231]:


def revNrep(string, sub1, sub2, count=1):
    if string.endswith(sub1):
        string = string[::-1]

        string = string.replace(sub1[::-1], sub2[::-1], count)

        string = string[::-1]
    return string
def removeEndings(string):
    y=string.replace('u','v')
    y=y.replace('j','i')
    
 
    y=revNrep(y,'int','')
    y=revNrep(y,'ent','')
    y=revNrep(y,'ient','')
    y=revNrep(y,'vnt','')
    y=revNrep(y,'ant','')
    y=revNrep(y,'are','')
    y=revNrep(y,'ere','')
    y=revNrep(y,'ire','')
    y=revNrep(y,'eat','')
    y=revNrep(y,'it','')
    y=revNrep(y,'et','')
    y=revNrep(y,'at','')
    y=revNrep(y,'atvr','')
    y=revNrep(y,'antvr','')

    
    y=revNrep(y,'arvm','')
    y=revNrep(y,'orvm','')
    y=revNrep(y,'ibvs','')
    y=revNrep(y,'vm','')
    y=revNrep(y,'vs','')
    y=revNrep(y,'es','')
    y=revNrep(y,'o','')
    y=revNrep(y,'am','')
    y=revNrep(y,'em','')
    y=revNrep(y,'is','')
    y=revNrep(y,'ias','')
    y=revNrep(y,'ia','')
    y=revNrep(y,'ae','')
    y=revNrep(y,'ii','')
    y=revNrep(y,'i','')
    y=revNrep(y,'e','')
    y=revNrep(y,'ate','')
    y=revNrep(y,'v','')
    y=revNrep(y,'a','')
    y=revNrep(y,'os','')
    
    return y

latinstem = []
for x in myWordList:  
    
    y = removeEndings(x)
    y = removeEndings(y)
    y = removeEndings(y)
    
    if not y:
        continue
    latinstem.append (y)
#print(latinstem)


# In[232]:


listForDict = list(set(latinstem))
listForDict.sort()
#print (listForDict) 


# In[233]:


for x in listForDict:
    #print(x)
    pass


# In[234]:


for i in range(10):
    mylen = len(list(filter(lambda x: len(x)==i,listForDict)))
    print(i,mylen)
mylen = len(list(filter(lambda x: len(x)>10,listForDict)))
print('>9',mylen)


# In[235]:


listForDict = list(filter(lambda x: len(x)>2,listForDict))
#print(listForDict)


# In[236]:


ListOfPrep = ['per','cum', 'si', 'set', 'dum', 'vel', 'sive', 'atqve', 'qve', 'eciam', 'omn', 'qvod', 'ad', 'in', 'sin', 'non', 'ut', 'tum','ill', 'ex', 'hvnc', 'ib', 
            'id', 'aut', 'hanc', 'hec', 'ilic', 'inter', 'ips', 'nec' , 'nil', 'neq', 'nichil', 'nul', 'nvnc', 'nvnq', 'qval', 'qvalemcvnq','qvaliter', 'qvamd', 'qvand',
            'qvandocvnq', 'qvas', 'qvas', 'qvatescvnq', 'qvemlib', 'qvicqvid', 'qvicvnq', 'qvid', 'qvilib', 'qvocvmq', 'qvocvnq', 'qvod', 'qvodcvnq', 'qvodlib', 'qvon',
            'qvoq', 'aliqvod', 'vnd', 'vsqveq', 'vsq', 'tot', 'tamen', 'svb', 'svper', 'sin', 'siqvid', 'sicvt', 'ess', 'swm', 'fuer', 'aput', 'svas', 'vbicvmq', 'vbicvnq',
             'vnoqvoq', 'vnvmqvodq', 'vnvsqvisq', 'sed', 'namq', 'illvd', 'illas', 'enim', 'eivsdemq', 'hoc', 'eqval', 'eqvaliter', 'aliqvand', 'aliqvid', 'aliqvod', 'ab']
ListOfPrepRem = []
for x in ListOfPrep:  
    y = removeEndings(x)
    ListOfPrepRem.append(y)
print(ListOfPrepRem)


# In[240]:



class DictLatCross():
    def __init__(self, ListOfPrep, listForDict, latinStem):
        self.ListOfPrep = ListOfPrep
        
        listForDict = list(filter(lambda x: not x in ListOfPrep, listForDict))
        
        latinStem = list(filter(lambda x: x in listForDict, latinStem))
        
        self.listForDict = listForDict
        self.latinStem = latinStem
        #print(self.listForDict)
        self.numeratedDict = dict(zip(self.listForDict, range(len(self.listForDict))))
        #print(self.latinStem)
        self.matrix = np.zeros((len(self.listForDict),len(self.listForDict)), dtype=np.int8)
        self.walkAlongText()
        
    def addToMatrix(self, word1, word2):
        self.matrix[self.numeratedDict[word1],self.numeratedDict[word2]]+=1
        self.matrix[self.numeratedDict[word2],self.numeratedDict[word1]]+=1
        #print(self.matrix[self.numeratedDict[word2],self.numeratedDict[word1]])
        
    def walkAlongText(self):
        for i in range(len(self.latinStem)-1):
            #print(self.latinStem[i],self.latinStem[i+1])
            #print(str(i)+'/'+str(len(self.latinStem)-1))
            self.addToMatrix(self.latinStem[i],self.latinStem[i+1])
            
    def findNMax(self, N = 10):
        #Выводит N самых частых пар
        def get_key(my_dict, val): 
            for key, value in my_dict.items(): 
                if val == value: 
                    return key   
        #N - число выводимых пар
        N = N * 2
        ind = np.argpartition(self.matrix.flatten(), -N)[-N:]
        ind = list(ind)
        ind.reverse()
        listOfPare = []

        for x in ind:
            a, b = unravel_index(x, self.matrix.shape)
            if a > b:
                a, b = b, a
            if not [a, b] in listOfPare:
                listOfPare.append([a,b])

        def sortMax(val):
            a, b = val
            return self.matrix[a,b]

        listOfPare.sort(key = sortMax, reverse = True)

        for x in listOfPare:
            a, b = x
            print(self.matrix[a,b], get_key(self.numeratedDict, a),get_key(self.numeratedDict, b))


# In[241]:


finalDict = DictLatCross(ListOfPrepRem, listForDict, latinstem)

finalDict.findNMax(20)


# In[239]:


len(finalDict.latinStem)


# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:




